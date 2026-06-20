"""
Generate Form Kuesioner TERISI per Responden — Audit SI PT. Murni Solusindo Nusantara
Menghasilkan 8 file DOCX (1 per responden) dengan jawaban sudah terisi (✓)
"""

import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "form_responden")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ============================================================
# DATA RESPONDEN & JAWABAN (sesuai Tabulasi_Data.md)
# ============================================================
RESPONDENTS = [
    {
        'code': 'R1', 'initial': 'AW', 'name': 'Ahmad Wijaya',
        'jabatan': 'Kepala Divisi IT', 'divisi': 'IT',
        'lama': '8 tahun', 'pendidikan': 'S1 Teknik Informatika',
        'tanggal': '15 / 03 / 2026',
    },
    {
        'code': 'R2', 'initial': 'BH', 'name': 'Budi Hartono',
        'jabatan': 'Senior Web Developer', 'divisi': 'Digital Marketing (Web Dev & SEO)',
        'lama': '5 tahun', 'pendidikan': 'S1 Sistem Informasi',
        'tanggal': '15 / 03 / 2026',
    },
    {
        'code': 'R3', 'initial': 'CS', 'name': 'Citra Sari',
        'jabatan': 'Junior Web Developer', 'divisi': 'Digital Marketing (Web Dev & SEO)',
        'lama': '2 tahun', 'pendidikan': 'S1 Sistem Informasi',
        'tanggal': '16 / 03 / 2026',
    },
    {
        'code': 'R4', 'initial': 'DK', 'name': 'Dian Kusuma',
        'jabatan': 'Content & SEO Specialist', 'divisi': 'Digital Marketing (Web Dev & SEO)',
        'lama': '4 tahun', 'pendidikan': 'S1 Teknik Informatika',
        'tanggal': '16 / 03 / 2026',
    },
    {
        'code': 'R5', 'initial': 'EP', 'name': 'Eko Prasetyo',
        'jabatan': 'Staff IT Support', 'divisi': 'IT',
        'lama': '3 tahun', 'pendidikan': 'D3 Teknik Komputer',
        'tanggal': '17 / 03 / 2026',
    },
    {
        'code': 'R6', 'initial': 'FR', 'name': 'Fajar Rahman',
        'jabatan': 'Digital Marketing Specialist', 'divisi': 'Digital Marketing',
        'lama': '3 tahun', 'pendidikan': 'S1 Ilmu Komunikasi',
        'tanggal': '17 / 03 / 2026',
    },
    {
        'code': 'R7', 'initial': 'GS', 'name': 'Galih Santoso',
        'jabatan': 'Staff Finance (Admin Sistem)', 'divisi': 'Finance',
        'lama': '4 tahun', 'pendidikan': 'S1 Akuntansi',
        'tanggal': '18 / 03 / 2026',
    },
    {
        'code': 'R8', 'initial': 'HL', 'name': 'Hesti Lestari',
        'jabatan': 'Staff HR (Kebijakan TI)', 'divisi': 'Human Resources',
        'lama': '5 tahun', 'pendidikan': 'S1 Manajemen SDM',
        'tanggal': '18 / 03 / 2026',
    },
]

# Jawaban per responden per domain (sesuai Tabulasi_Data.md)
# Format: ANSWERS[domain_idx][question_idx] = [R1, R2, ..., R8]
ANSWERS = {
    'APO12': [
        # Level 1
        [4, 4, 3, 4, 3, 3, 3, 3],  # Q1
        [4, 3, 3, 4, 3, 3, 3, 2],  # Q2
        [3, 3, 3, 3, 2, 3, 2, 2],  # Q3
        # Level 2
        [4, 3, 3, 3, 3, 3, 2, 2],  # Q4
        [3, 3, 2, 3, 2, 3, 2, 2],  # Q5
        [3, 3, 2, 3, 2, 2, 2, 2],  # Q6
        # Level 3
        [3, 3, 2, 3, 2, 2, 2, 2],  # Q7
        [3, 2, 2, 3, 2, 2, 2, 2],  # Q8
        [3, 2, 2, 2, 2, 2, 2, 2],  # Q9
    ],
    'APO13': [
        [4, 4, 3, 4, 3, 3, 3, 3],
        [4, 4, 4, 4, 3, 3, 3, 3],
        [3, 3, 3, 3, 3, 3, 3, 2],
        [4, 3, 3, 4, 3, 3, 3, 2],
        [3, 3, 3, 3, 2, 3, 2, 2],
        [3, 3, 2, 3, 2, 2, 2, 2],
        [3, 3, 2, 3, 2, 2, 2, 2],
        [3, 3, 3, 3, 2, 2, 2, 2],
        [2, 2, 2, 3, 2, 2, 2, 2],
    ],
    'BAI06': [
        [4, 4, 3, 4, 3, 3, 3, 2],
        [4, 4, 3, 3, 3, 3, 3, 2],
        [4, 3, 3, 3, 3, 3, 3, 2],
        [4, 3, 3, 3, 3, 3, 3, 2],
        [3, 3, 3, 3, 2, 2, 2, 2],
        [3, 3, 3, 3, 2, 2, 2, 2],
        [3, 3, 2, 3, 2, 2, 2, 2],
        [3, 3, 2, 3, 2, 2, 2, 2],
        [3, 2, 2, 3, 2, 2, 2, 2],
    ],
    'DSS01': [
        [4, 4, 4, 4, 3, 3, 3, 3],
        [4, 3, 3, 4, 3, 3, 3, 3],
        [4, 4, 3, 4, 3, 3, 3, 2],
        [4, 3, 3, 4, 3, 3, 3, 2],
        [3, 3, 3, 3, 3, 3, 2, 2],
        [3, 3, 3, 3, 2, 2, 2, 2],
        [3, 3, 3, 3, 2, 2, 2, 2],
        [3, 3, 2, 3, 2, 2, 2, 2],
        [3, 3, 2, 3, 2, 2, 2, 2],
    ],
    'DSS05': [
        [4, 4, 3, 4, 3, 3, 3, 3],
        [4, 4, 4, 4, 3, 3, 3, 3],
        [4, 4, 3, 4, 3, 3, 3, 2],
        [4, 3, 3, 3, 3, 3, 2, 2],
        [3, 3, 3, 3, 2, 2, 2, 2],
        [3, 3, 2, 3, 2, 2, 2, 2],
        [3, 3, 2, 3, 2, 2, 2, 2],
        [3, 2, 2, 3, 2, 2, 2, 2],
        [2, 2, 2, 3, 2, 2, 2, 2],
    ],
}

# Domain descriptions + questions (same as form template)
DOMAINS = [
    {
        'code': 'APO12',
        'name': 'Managed Risk (Pengelolaan Risiko)',
        'desc': 'Domain ini mengevaluasi sejauh mana organisasi telah mengelola risiko teknologi informasi secara sistematis.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Risiko TI yang terkait dengan platform digital telah diidentifikasi secara berkala',
                'Pengumpulan data mengenai risiko TI dilakukan secara efektif',
                'Dokumentasi risiko TI dipelihara dan diperbarui secara teratur',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan untuk pengelolaan risiko TI',
                'Risiko TI dipantau dan dilaporkan kepada pihak terkait secara berkala',
                'Dilakukan evaluasi terhadap pengelolaan risiko secara rutin',
            ],
            'Level 3 (Established Process)': [
                'Terdapat prosedur standar yang terdokumentasi untuk mengelola risiko TI',
                'Prosedur pengelolaan risiko diterapkan secara konsisten di seluruh divisi',
                'Dilakukan audit internal terhadap pengelolaan risiko secara periodik',
            ],
        },
    },
    {
        'code': 'APO13',
        'name': 'Managed Security (Pengelolaan Keamanan)',
        'desc': 'Domain ini mengevaluasi kebijakan dan pengelolaan keamanan informasi organisasi.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Kebijakan keamanan informasi telah ditetapkan dan dikomunikasikan',
                'Mekanisme perlindungan terhadap akses tidak sah telah diterapkan',
                'Sosialisasi mengenai keamanan informasi sudah dilakukan kepada karyawan',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan untuk pengelolaan keamanan informasi',
                'Insiden keamanan informasi dipantau dan ditindaklanjuti',
                'Evaluasi terhadap kebijakan keamanan dilakukan secara berkala',
            ],
            'Level 3 (Established Process)': [
                'Terdapat prosedur standar keamanan informasi yang terdokumentasi',
                'Prosedur keamanan diterapkan secara konsisten di seluruh platform digital',
                'Dilakukan pelatihan keamanan informasi secara rutin',
            ],
        },
    },
    {
        'code': 'BAI06',
        'name': 'Managed IT Changes (Pengelolaan Perubahan TI)',
        'desc': 'Domain ini mengevaluasi pengelolaan perubahan pada sistem dan infrastruktur TI.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Setiap perubahan pada sistem TI dicatat dan didokumentasikan',
                'Terdapat proses persetujuan sebelum implementasi perubahan',
                'Evaluasi dampak perubahan dilakukan sebelum implementasi',
            ],
            'Level 2 (Managed Process)': [
                'Perubahan TI direncanakan dan dijadwalkan secara terstruktur',
                'Terdapat komunikasi mengenai perubahan kepada pihak terkait',
                'Evaluasi pasca-implementasi perubahan dilakukan secara rutin',
            ],
            'Level 3 (Established Process)': [
                'Terdapat prosedur standar change management yang terdokumentasi',
                'Prosedur change management diterapkan konsisten untuk semua perubahan',
                'Dilakukan review berkala terhadap efektivitas prosedur change management',
            ],
        },
    },
    {
        'code': 'DSS01',
        'name': 'Managed Operations (Pengelolaan Operasional)',
        'desc': 'Domain ini mengevaluasi pengelolaan operasional harian teknologi informasi.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Pemantauan terhadap kinerja platform digital dilakukan secara rutin',
                'Terdapat prosedur penanganan insiden operasional',
                'Jadwal pemeliharaan infrastruktur TI sudah ditetapkan',
            ],
            'Level 2 (Managed Process)': [
                'Operasional TI direncanakan dan dikelola secara terstruktur',
                'Kinerja operasional TI dipantau dan dilaporkan secara berkala',
                'Evaluasi terhadap pengelolaan operasional dilakukan rutin',
            ],
            'Level 3 (Established Process)': [
                'Terdapat SOP operasional TI yang terdokumentasi secara formal',
                'SOP operasional diterapkan secara konsisten di seluruh divisi',
                'Terdapat mekanisme feedback dari pengguna terhadap layanan TI',
            ],
        },
    },
    {
        'code': 'DSS05',
        'name': 'Managed Security Services (Pengelolaan Layanan Keamanan)',
        'desc': 'Domain ini mengevaluasi implementasi teknis layanan keamanan TI.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perlindungan terhadap malware dan ancaman siber telah diterapkan',
                'Pengelolaan hak akses pengguna (access control) sudah dilaksanakan',
                'Prosedur pencadangan dan pemulihan data sudah berjalan',
            ],
            'Level 2 (Managed Process)': [
                'Layanan keamanan TI direncanakan dan dikelola secara terstruktur',
                'Insiden keamanan TI dipantau dan ditindaklanjuti secara berkala',
                'Evaluasi efektivitas layanan keamanan dilakukan rutin',
            ],
            'Level 3 (Established Process)': [
                'Terdapat prosedur standar pengelolaan layanan keamanan yang terdokumentasi',
                'Standar keamanan diterapkan konsisten di seluruh platform digital',
                'Audit keamanan internal dilaksanakan secara periodik',
            ],
        },
    },
]


# ============================================================
# GENERATE PER RESPONDEN
# ============================================================
def generate_form(resp_idx, resp):
    """Generate satu DOCX form terisi untuk satu responden"""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)

    def centered(text, size=12, bold=False, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold

    def normal(text, bold=False, indent=False, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(space_after)
        if indent:
            p.paragraph_format.first_line_indent = Cm(1.25)
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = bold

    def set_cell_font(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, bg=None):
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = align
        for r in p.runs:
            r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold
            if color:
                r.font.color.rgb = color
        if bg:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

    # ── COVER ──
    centered("KUESIONER PENELITIAN", 16, bold=True, space_after=12)
    centered("AUDIT SISTEM INFORMASI", 14, bold=True, space_after=6)
    centered("PT. MURNI SOLUSINDO NUSANTARA", 14, bold=True, space_after=6)
    centered("Framework COBIT 2019", 12, space_after=18)
    centered(f"— Form Isian Responden {resp['code']} —", 12, bold=True, space_after=24)

    # ── IDENTITAS RESPONDEN (TERISI) ──
    centered("IDENTITAS RESPONDEN", 14, bold=True, space_after=12)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    fields = [
        ("Nama", resp['name']),
        ("Jabatan", resp['jabatan']),
        ("Divisi / Departemen", resp['divisi']),
        ("Lama Bekerja", resp['lama']),
        ("Pendidikan Terakhir", resp['pendidikan']),
        ("Tanggal Pengisian", resp['tanggal']),
    ]
    for i, (label, value) in enumerate(fields):
        set_cell_font(table.rows[i].cells[0], label, 11)
        set_cell_font(table.rows[i].cells[1], ":", 11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_font(table.rows[i].cells[2], value, 11, bold=True)
        table.rows[i].cells[0].width = Cm(5)
        table.rows[i].cells[1].width = Cm(0.5)
        table.rows[i].cells[2].width = Cm(9)

    doc.add_page_break()

    # ── KUESIONER PER DOMAIN (JAWABAN TERISI) ──
    q_global = 0
    for d_idx, domain in enumerate(DOMAINS):
        centered(f"DOMAIN {domain['code']}", 14, bold=True, space_after=2)
        centered(domain['name'], 12, bold=True, space_after=6)
        normal(domain['desc'], space_after=12)

        for level_name, questions in domain['levels'].items():
            normal(level_name, bold=True, space_after=6)

            num_q = len(questions)
            table = doc.add_table(rows=1 + num_q, cols=7)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Header
            for j, h in enumerate(["No", "Pernyataan", "STS (1)", "TS (2)", "R (3)", "S (4)", "SS (5)"]):
                set_cell_font(table.rows[0].cells[j], h, 10, bold=True,
                             align=WD_ALIGN_PARAGRAPH.CENTER,
                             color=RGBColor(255, 255, 255), bg="1B3A5C")

            # Questions with filled answers
            for qi, question in enumerate(questions):
                row = table.rows[qi + 1]
                q_num = q_global + 1

                set_cell_font(row.cells[0], str(q_num), 10, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_font(row.cells[1], question, 10)

                # Get this respondent's answer
                answer = ANSWERS[domain['code']][q_global % 9][resp_idx]

                # Fill checkmark on the correct column
                for ci in range(2, 7):
                    score = ci - 1  # col 2=1(STS), col 3=2(TS), col 4=3(R), col 5=4(S), col 6=5(SS)
                    if score == answer:
                        set_cell_font(row.cells[ci], "✓", 12, bold=True,
                                     align=WD_ALIGN_PARAGRAPH.CENTER)
                    else:
                        row.cells[ci].text = ""
                        row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row.cells[ci].width = Cm(1.5)

                row.cells[0].width = Cm(1)
                row.cells[1].width = Cm(8)
                q_global += 1

            doc.add_paragraph().paragraph_format.space_after = Pt(8)

        if d_idx < len(DOMAINS) - 1:
            doc.add_page_break()

    # ── FOOTER ──
    doc.add_paragraph()
    normal("— Terima kasih atas kesediaan Bapak/Ibu mengisi kuesioner ini —", bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("Tanda tangan responden:"); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"( {resp['name']} )"); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.bold = True

    # Save
    filename = f"Kuesioner_Terisi_{resp['code']}_{resp['initial']}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filename


# ============================================================
# MAIN
# ============================================================
print("=" * 60)
print("  GENERATE FORM KUESIONER TERISI PER RESPONDEN")
print("=" * 60)

files = []
for idx, resp in enumerate(RESPONDENTS):
    fname = generate_form(idx, resp)
    files.append(fname)
    print(f"  ✅ {resp['code']} ({resp['initial']}) — {resp['jabatan']:30s} → {fname}")

print(f"\n📂 Output folder: {OUTPUT_DIR}")
print(f"📄 Total file: {len(files)} form kuesioner terisi")
print(f"📝 Setiap form berisi: Identitas (terisi) + 45 soal (jawaban ✓)")
