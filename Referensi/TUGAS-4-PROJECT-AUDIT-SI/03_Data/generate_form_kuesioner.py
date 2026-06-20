"""
Generate Form Kuesioner Word (.docx) — Audit SI PT. Murni Solusindo Nusantara
Form ini bisa dicetak atau dikirim ke responden untuk diisi
"""

import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "Form_Kuesioner_Audit_SI.docx")

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)


def centered(text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold
    return p

def normal(text, bold=False, indent=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = bold
    return p


# ============================================================
# HALAMAN 1: COVER & IDENTITAS
# ============================================================
centered("KUESIONER PENELITIAN", 16, bold=True, space_after=12)
centered("AUDIT SISTEM INFORMASI", 14, bold=True, space_after=6)
centered("PT. MURNI SOLUSINDO NUSANTARA", 14, bold=True, space_after=18)
centered("Framework COBIT 2019", 12, space_after=24)

normal("Kepada Yth. Bapak/Ibu Responden", bold=True, space_after=6)
normal("di Tempat", space_after=12)

normal("Dengan hormat,", indent=True, space_after=6)
normal("Perkenalkan, kami mahasiswa Program Studi Sistem Informasi (S1), Fakultas Teknologi Informasi, Universitas Nusa Mandiri yang sedang melakukan penelitian dalam rangka memenuhi tugas mata kuliah Audit Sistem Informasi. Penelitian ini bertujuan untuk mengevaluasi tata kelola teknologi informasi pada PT. Murni Solusindo Nusantara menggunakan framework COBIT 2019.", indent=True, space_after=6)
normal("Kami mengharapkan kesediaan Bapak/Ibu untuk mengisi kuesioner ini dengan sebaik-baiknya. Data yang diperoleh akan digunakan semata-mata untuk kepentingan akademis dan dijamin kerahasiaannya.", indent=True, space_after=6)
normal("Atas perhatian dan kesediaan Bapak/Ibu, kami ucapkan terima kasih.", indent=True, space_after=12)

# Tim peneliti
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(12)
r = p.add_run("Hormat kami,"); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Tim Peneliti"); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True

doc.add_page_break()


# ============================================================
# HALAMAN 2: IDENTITAS RESPONDEN
# ============================================================
centered("IDENTITAS RESPONDEN", 14, bold=True, space_after=18)

normal("Petunjuk: Mohon lengkapi data berikut.", space_after=12)

# Identity table
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
fields = [
    ("Nama", ""),
    ("Jabatan", ""),
    ("Divisi / Departemen", ""),
    ("Lama Bekerja", "...... tahun"),
    ("Pendidikan Terakhir", ""),
    ("Tanggal Pengisian", "...... / ...... / 2026"),
]
for i, (label, default) in enumerate(fields):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = ":"
    table.rows[i].cells[2].text = default
    for j in range(3):
        p = table.rows[i].cells[j].paragraphs[0]
        p.runs[0].font.name = 'Times New Roman' if p.runs else None
        p.runs[0].font.size = Pt(12) if p.runs else None
    table.rows[i].cells[0].width = Cm(5)
    table.rows[i].cells[1].width = Cm(0.5)
    table.rows[i].cells[2].width = Cm(9)

doc.add_paragraph().paragraph_format.space_after = Pt(18)

# Petunjuk pengisian
centered("PETUNJUK PENGISIAN", 14, bold=True, space_after=12)
normal("Berilah tanda centang (✓) pada salah satu kolom yang sesuai dengan pendapat Bapak/Ibu terhadap pernyataan-pernyataan berikut:", space_after=12)

# Scale explanation
scale_table = doc.add_table(rows=6, cols=3)
scale_table.style = 'Table Grid'
scale_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header
for j, h in enumerate(["Skor", "Singkatan", "Keterangan"]):
    cell = scale_table.rows[0].cells[j]
    cell.text = h
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.bold = True
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    for r in p.runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

scales = [
    ("1", "STS", "Sangat Tidak Setuju"),
    ("2", "TS", "Tidak Setuju"),
    ("3", "R", "Ragu-ragu / Netral"),
    ("4", "S", "Setuju"),
    ("5", "SS", "Sangat Setuju"),
]
for i, (skor, singkat, ket) in enumerate(scales):
    for j, val in enumerate([skor, singkat, ket]):
        cell = scale_table.rows[i+1].cells[j]
        cell.text = val
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'; r.font.size = Pt(11)

doc.add_page_break()


# ============================================================
# HALAMAN 3+: KUESIONER PER DOMAIN
# ============================================================
domains = [
    {
        'code': 'APO12',
        'name': 'Managed Risk (Pengelolaan Risiko)',
        'desc': 'Domain ini mengevaluasi sejauh mana organisasi telah mengelola risiko teknologi informasi secara sistematis.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Risiko TI yang terkait dengan platform digital telah diidentifikasi secara berkala',
                'Pengumpulan data mengenai risiko TI dilakukan secara efektif',
                'Dokumentasi risiko TI dipelihara dan diperbarui secara teratur',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan untuk pengelolaan risiko TI',
                'Risiko TI dipantau dan dilaporkan kepada pihak terkait secara berkala',
                'Dilakukan evaluasi terhadap pengelolaan risiko secara rutin',
            ],
            'Level 3 (Established Process)': [
                'Terdapat prosedur standar yang terdokumentasi untuk mengelola risiko TI',
                'Prosedur pengelolaan risiko diterapkan secara konsisten di seluruh divisi',
                'Dilakukan audit internal terhadap pengelolaan risiko secara periodik',
            ],
        },
    },
    {
        'code': 'APO13',
        'name': 'Managed Security (Pengelolaan Keamanan)',
        'desc': 'Domain ini mengevaluasi kebijakan dan pengelolaan keamanan informasi organisasi.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Kebijakan keamanan informasi telah ditetapkan dan dikomunikasikan',
                'Mekanisme perlindungan terhadap akses tidak sah telah diterapkan',
                'Sosialisasi mengenai keamanan informasi sudah dilakukan kepada karyawan',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan untuk pengelolaan keamanan informasi',
                'Insiden keamanan informasi dipantau dan ditindaklanjuti',
                'Evaluasi terhadap kebijakan keamanan dilakukan secara berkala',
            ],
            'Level 3 (Established Process)': [
                'Terdapat prosedur standar keamanan informasi yang terdokumentasi',
                'Prosedur keamanan diterapkan secara konsisten di seluruh platform digital',
                'Dilakukan pelatihan keamanan informasi secara rutin',
            ],
        },
    },
    {
        'code': 'BAI06',
        'name': 'Managed IT Changes (Pengelolaan Perubahan TI)',
        'desc': 'Domain ini mengevaluasi pengelolaan perubahan pada sistem dan infrastruktur TI.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Setiap perubahan pada sistem TI dicatat dan didokumentasikan',
                'Terdapat proses persetujuan sebelum implementasi perubahan',
                'Evaluasi dampak perubahan dilakukan sebelum implementasi',
            ],
            'Level 2 (Managed Process)': [
                'Perubahan TI direncanakan dan dijadwalkan secara terstruktur',
                'Terdapat komunikasi mengenai perubahan kepada pihak terkait',
                'Evaluasi pasca-implementasi perubahan dilakukan secara rutin',
            ],
            'Level 3 (Established Process)': [
                'Terdapat prosedur standar change management yang terdokumentasi',
                'Prosedur change management diterapkan konsisten untuk semua perubahan',
                'Dilakukan review berkala terhadap efektivitas prosedur change management',
            ],
        },
    },
    {
        'code': 'DSS01',
        'name': 'Managed Operations (Pengelolaan Operasional)',
        'desc': 'Domain ini mengevaluasi pengelolaan operasional harian teknologi informasi.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Pemantauan terhadap kinerja platform digital dilakukan secara rutin',
                'Terdapat prosedur penanganan insiden operasional',
                'Jadwal pemeliharaan infrastruktur TI sudah ditetapkan',
            ],
            'Level 2 (Managed Process)': [
                'Operasional TI direncanakan dan dikelola secara terstruktur',
                'Kinerja operasional TI dipantau dan dilaporkan secara berkala',
                'Evaluasi terhadap pengelolaan operasional dilakukan rutin',
            ],
            'Level 3 (Established Process)': [
                'Terdapat SOP operasional TI yang terdokumentasi secara formal',
                'SOP operasional diterapkan secara konsisten di seluruh divisi',
                'Terdapat mekanisme feedback dari pengguna terhadap layanan TI',
            ],
        },
    },
    {
        'code': 'DSS05',
        'name': 'Managed Security Services (Pengelolaan Layanan Keamanan)',
        'desc': 'Domain ini mengevaluasi implementasi teknis layanan keamanan TI.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perlindungan terhadap malware dan ancaman siber telah diterapkan',
                'Pengelolaan hak akses pengguna (access control) sudah dilaksanakan',
                'Prosedur pencadangan dan pemulihan data sudah berjalan',
            ],
            'Level 2 (Managed Process)': [
                'Layanan keamanan TI direncanakan dan dikelola secara terstruktur',
                'Insiden keamanan TI dipantau dan ditindaklanjuti secara berkala',
                'Evaluasi efektivitas layanan keamanan dilakukan rutin',
            ],
            'Level 3 (Established Process)': [
                'Terdapat prosedur standar pengelolaan layanan keamanan yang terdokumentasi',
                'Standar keamanan diterapkan konsisten di seluruh platform digital',
                'Audit keamanan internal dilaksanakan secara periodik',
            ],
        },
    },
]

q_global = 1
for d_idx, domain in enumerate(domains):
    centered(f"DOMAIN {domain['code']}", 14, bold=True, space_after=2)
    centered(domain['name'], 12, bold=True, space_after=6)
    normal(domain['desc'], space_after=12)

    for level_name, questions in domain['levels'].items():
        normal(level_name, bold=True, space_after=6)

        # Create table for questions
        num_q = len(questions)
        table = doc.add_table(rows=1 + num_q, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        for j, h in enumerate(["No", "Pernyataan", "STS (1)", "TS (2)", "R (3)", "S (4)", "SS (5)"]):
            cell = table.rows[0].cells[j]
            cell.text = h
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # Questions
        for qi, question in enumerate(questions):
            row = table.rows[qi + 1]
            # No
            row.cells[0].text = str(q_global)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in row.cells[0].paragraphs[0].runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)

            # Question text
            row.cells[1].text = question
            for r in row.cells[1].paragraphs[0].runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)

            # Empty checkboxes (STS-SS)
            for ci in range(2, 7):
                row.cells[ci].text = ""
                row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set fixed width for checkbox columns
                row.cells[ci].width = Cm(1.5)

            # Set column widths
            row.cells[0].width = Cm(1)
            row.cells[1].width = Cm(8)

            q_global += 1

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    if d_idx < len(domains) - 1:
        doc.add_page_break()

# Footer note
doc.add_paragraph()
normal("— Terima kasih atas kesediaan Bapak/Ibu mengisi kuesioner ini —", bold=True, space_after=6)
normal("Data yang Bapak/Ibu berikan akan dijaga kerahasiaannya dan hanya digunakan untuk kepentingan akademis.", space_after=6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(24)
r = p.add_run("Tanda tangan responden:"); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("(__________________________)"); r.font.name = 'Times New Roman'; r.font.size = Pt(11)


# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
print(f"✅ Form Kuesioner berhasil disimpan: {OUTPUT}")
print(f"📋 Total domain: 5 (APO12, APO13, BAI06, DSS01, DSS05)")
print(f"📝 Total pertanyaan: {q_global - 1} soal (9 per domain × 5 domain = 45)")
print(f"📄 Isi: Cover + Identitas Responden + Petunjuk + Kuesioner 5 Domain")
