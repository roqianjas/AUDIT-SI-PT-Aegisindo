"""
Konversi Panduan & Contekan Presentasi (Markdown -> Word .docx)
Format: TNR 12pt, Spasi 1.15, Margin 2.5 cm.
Mempertahankan blockquote (naskah ngomong) sebagai kotak terindentasi.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === PATHS ===
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SOURCE_PATH = os.path.join(BASE_DIR, "PANDUAN_PRESENTASI.md")
LOGO_PATH = os.path.join(BASE_DIR, "assets", "images", "logo-unm.png")
OUTPUT_PATH = os.path.join(BASE_DIR, "Kelompok-Panduan-Presentasi-Audit-SI-PT_Murni.docx")

FONT = "Times New Roman"

# Warna penanda pembicara -> nama (untuk dihilangkan dari judul + diberi label)
SPEAKER_BY_EMOJI = {
    "🟦": "Roki",
    "🟩": "Susanto",
}
# Warna kotak penanda per pembicara (RGB)
SPEAKER_COLOR = {
    "Roki": RGBColor(0x25, 0x63, 0xEB),       # biru
    "Susanto": RGBColor(0x16, 0xA3, 0x4A),    # hijau
}
EMOJI_RE = re.compile(r"[🟦🟩🟨🟪🎤👥📋⏱️🧠📑❓📖✅💪🟥]")


def read_file(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


# === SETUP DOCUMENT ===
doc = Document()

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

for i, size in [(1, 16), (2, 13), (3, 12)]:
    hs = doc.styles[f"Heading {i}"]
    hs.font.name = FONT
    hs.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    hs.font.bold = True
    hs.font.size = Pt(size)
    hs.paragraph_format.line_spacing = 1.15
    hs.paragraph_format.space_before = Pt(10)
    hs.paragraph_format.space_after = Pt(4)


def strip_md(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"<br\s*/?>", "; ", text)
    return text.strip()


def add_runs_with_bold(paragraph, text):
    """Render **bold** spans as bold runs, sisanya normal."""
    text = re.sub(r"`(.+?)`", r"\1", text)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        else:
            part = re.sub(r"\*(.+?)\*", r"\1", part)
            run = paragraph.add_run(part)
        run.font.name = FONT
        run.font.size = Pt(12)


def shade_paragraph(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_quote(text):
    """Naskah ngomong: indentasi + shading abu muda + italic."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.right_indent = Cm(0.75)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    shade_paragraph(p, "F1F1F5")
    run = p.add_run(strip_md(text.lstrip(">").strip().strip('"')))
    run.font.name = FONT
    run.font.size = Pt(12)
    run.font.italic = True
    return p


def add_para(text, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(space_after)
    add_runs_with_bold(p, text)
    return p


def add_list_item(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    add_runs_with_bold(p, text)
    return p


def add_centered(text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    return p


def parse_md_table(block_lines):
    lines = [l for l in block_lines if not re.match(r"^\|[-\s|:]+\|$", l.strip())]
    if not lines:
        return [], []
    headers = [c.strip() for c in lines[0].split("|")[1:-1]]
    rows = []
    for line in lines[1:]:
        cols = [c.strip() for c in line.split("|")[1:-1]]
        if cols:
            rows.append(cols)
    return headers, rows


def add_table(headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tbl.tblPr.append(tblW)

    hdr = table.add_row()
    for i, h in enumerate(headers):
        cp = hdr.cells[i].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(strip_md(h))
        r.font.name = FONT
        r.font.size = Pt(10)
        r.font.bold = True
        cp.paragraph_format.line_spacing = 1.0
        cp.paragraph_format.space_after = Pt(2)
        cp.paragraph_format.space_before = Pt(2)
        shade_paragraph(cp, "E0E7FF")

    for row_data in rows:
        row = table.add_row()
        for ci, val in enumerate(row_data):
            if ci >= len(headers):
                break
            cp = row.cells[ci].paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = cp.add_run(strip_md(str(val)))
            r.font.name = FONT
            r.font.size = Pt(10)
            cp.paragraph_format.line_spacing = 1.0
            cp.paragraph_format.space_after = Pt(2)
            cp.paragraph_format.space_before = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C7D2FE")
    pbdr.append(bottom)
    pPr.append(pbdr)


# === COVER ===
add_centered("PANDUAN & CONTEKAN PRESENTASI", 16, bold=True, space_after=2)
add_centered("Audit Tata Kelola TI PT. Murni Solusindo Nusantara", 13, bold=True, space_after=2)
add_centered("Berbasis COBIT 2019", 13, bold=True, space_after=14)

if os.path.exists(LOGO_PATH):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(14)
    p_logo.add_run().add_picture(LOGO_PATH, width=Cm(3.5), height=Cm(3.5))

add_centered("Kelompok — Program Studi Sistem Informasi (S1)", 12, bold=True, space_after=2)
add_centered("Mata Kuliah: Audit Sistem Informasi (453)", 12, space_after=2)
add_centered("Dosen Pengampu: Juarni Siregar, S.Pd., M.Kom", 12, space_after=2)
add_centered("Universitas Nusa Mandiri, Jakarta — 2026", 12, space_after=2)

doc.add_page_break()

# === PARSE MARKDOWN BODY ===
md = read_file(SOURCE_PATH)
lines = md.split("\n")
i = 0
n = len(lines)

# Lewati judul cover markdown (4 baris pertama: # judul, ## sub, ** **, ** **)
# supaya tidak duplikat dengan cover yang sudah dibuat.
while i < n and not lines[i].startswith("> Dokumen ini"):
    i += 1

while i < n:
    raw = lines[i].rstrip()
    line = raw.strip()

    if not line:
        i += 1
        continue

    if line == "---":
        add_hr()
        i += 1
        continue

    # Heading
    m = re.match(r"^(#{1,3})\s+(.+)", line)
    if m:
        level = len(m.group(1))
        title = strip_md(EMOJI_RE.sub("", m.group(2)).strip())
        emoji = next((e for e in SPEAKER_BY_EMOJI if e in m.group(2)), None)
        if emoji and title.upper().startswith("SLIDE"):
            speaker = SPEAKER_BY_EMOJI[emoji]
            h = doc.add_heading("", level=min(level, 3))
            box = h.add_run("■ ")  # kotak penuh berwarna
            box.font.name = FONT
            box.font.color.rgb = SPEAKER_COLOR[speaker]
            h.add_run(f"{title}   [{speaker}]")
        else:
            doc.add_heading(title, level=min(level, 3))
        i += 1
        continue

    # Blockquote (naskah ngomong) — bisa multi-baris
    if line.startswith(">"):
        quote_lines = []
        while i < n and lines[i].strip().startswith(">"):
            quote_lines.append(lines[i].strip().lstrip(">").strip())
            i += 1
        add_quote(" ".join(quote_lines))
        continue

    # Table
    if line.startswith("|"):
        block = []
        while i < n and lines[i].strip().startswith("|"):
            block.append(lines[i].strip())
            i += 1
        headers, rows = parse_md_table(block)
        if headers and rows:
            add_table(headers, rows)
        continue

    # Numbered list
    if re.match(r"^\d+\.\s+", line):
        add_list_item(line)
        i += 1
        continue

    # Bullet list (termasuk checklist - [ ])
    if line.startswith("- "):
        text = line[2:].strip()
        text = re.sub(r"^\[\s?\]\s*", "", text)  # buang checkbox markdown
        add_list_item(f"•  {text}")
        i += 1
        continue

    # Paragraf biasa
    add_para(line)
    i += 1

doc.save(OUTPUT_PATH)
print(f"[OK] Panduan presentasi tersimpan: {OUTPUT_PATH}")
print("[i] Format: TNR 12pt, spasi 1.15, margin 2.5 cm")
print("[i] Naskah ngomong (blockquote) dipertahankan sebagai kotak terindentasi")
