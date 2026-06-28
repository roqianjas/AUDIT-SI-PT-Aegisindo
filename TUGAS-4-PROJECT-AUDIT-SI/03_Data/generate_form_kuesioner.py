"""
Generate Form Kuesioner Word (.docx) — Audit SI PT Aegisindo Mitra Sejati
Form kosong yang bisa dicetak atau dikirim ke responden untuk diisi.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(BASE_DIR, "Form_Kuesioner_Audit_SI.docx")

HEADER_FILL = "F26522"  # safety orange

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)


def centered(text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold
    return p

def normal(text, bold=False, indent=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = bold
    return p


# ============================================================
# HALAMAN 1: COVER & PENGANTAR
# ============================================================
centered("KUESIONER PENELITIAN", 16, bold=True, space_after=12)
centered("AUDIT SISTEM INFORMASI", 14, bold=True, space_after=6)
centered("PT AEGISINDO MITRA SEJATI", 14, bold=True, space_after=18)
centered("Framework COBIT 2019", 12, space_after=24)

normal("Kepada Yth. Bapak/Ibu Responden", bold=True, space_after=6)
normal("di Tempat", space_after=12)

normal("Dengan hormat,", indent=True, space_after=6)
normal("Perkenalkan, kami mahasiswa Program Studi Sistem Informasi (S1), Fakultas Teknologi Informasi, Universitas Nusa Mandiri yang sedang melakukan penelitian dalam rangka memenuhi tugas mata kuliah Audit Sistem Informasi. Penelitian ini bertujuan mengaudit tata kelola teknologi informasi pada PT Aegisindo Mitra Sejati menggunakan framework COBIT 2019.", indent=True, space_after=6)
normal("Kami mengharapkan kesediaan Bapak/Ibu untuk mengisi kuesioner ini sesuai dengan kondisi nyata di perusahaan. Data yang diperoleh akan digunakan semata-mata untuk kepentingan akademis dan dijamin kerahasiaannya.", indent=True, space_after=6)
normal("Atas perhatian dan kesediaan Bapak/Ibu, kami ucapkan terima kasih.", indent=True, space_after=12)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(12)
r = p.add_run("Hormat kami,"); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Tim Peneliti"); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True

doc.add_page_break()


# ============================================================
# HALAMAN 2: IDENTITAS + PETUNJUK
# ============================================================
centered("IDENTITAS RESPONDEN", 14, bold=True, space_after=18)
normal("Petunjuk: Mohon lengkapi data berikut.", space_after=12)

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
fields = [
    ("Nama", ""),
    ("Jabatan", ""),
    ("Divisi / Bagian", ""),
    ("Lama Bekerja", "...... tahun"),
    ("Pendidikan Terakhir", ""),
    ("Tanggal Pengisian", "...... / ...... / 2026"),
]
for i, (label, default) in enumerate(fields):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = ":"
    table.rows[i].cells[2].text = default
    for j in range(3):
        cp = table.rows[i].cells[j].paragraphs[0]
        if cp.runs:
            cp.runs[0].font.name = 'Times New Roman'
            cp.runs[0].font.size = Pt(12)
    table.rows[i].cells[0].width = Cm(5)
    table.rows[i].cells[1].width = Cm(0.5)
    table.rows[i].cells[2].width = Cm(9)

doc.add_paragraph().paragraph_format.space_after = Pt(18)

centered("PETUNJUK PENGISIAN", 14, bold=True, space_after=12)
normal("Berilah tanda centang (✓) pada salah satu kolom yang sesuai dengan kondisi nyata di perusahaan terhadap pernyataan-pernyataan berikut:", space_after=12)

scale_table = doc.add_table(rows=6, cols=3)
scale_table.style = 'Table Grid'
scale_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, h in enumerate(["Skor", "Singkatan", "Keterangan"]):
    cell = scale_table.rows[0].cells[j]
    cell.text = h
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_FILL}"/>'))

scales = [
    ("1", "STS", "Sangat Tidak Setuju"),
    ("2", "TS", "Tidak Setuju"),
    ("3", "R", "Ragu-ragu / Netral"),
    ("4", "S", "Setuju"),
    ("5", "SS", "Sangat Setuju"),
]
for i, row_vals in enumerate(scales):
    for j, val in enumerate(row_vals):
        cell = scale_table.rows[i+1].cells[j]
        cell.text = val
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'; r.font.size = Pt(11)

doc.add_page_break()


# ============================================================
# HALAMAN 3+: KUESIONER PER DOMAIN
# ============================================================
domains = [
    {
        'code': 'APO14',
        'name': 'Managed Data (Pengelolaan Data)',
        'desc': 'Domain ini mengevaluasi sejauh mana perusahaan mengelola data inventori, pelanggan, dan penjualan secara terstruktur.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perusahaan telah mendata dan mengelola data inventori, pelanggan, dan penjualan yang dibutuhkan operasional',
                'Data produk dan stok dicatat dan diperbarui saat terjadi transaksi pembelian maupun penjualan',
                'Perusahaan menyimpan data pelanggan dan riwayat transaksi pada media penyimpanan tertentu',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan yang jelas mengenai data apa saja yang dikelola dan siapa penanggung jawabnya',
                'Kualitas dan konsistensi data antar saluran penjualan dipantau secara berkala',
                'Hasil pengelolaan data dievaluasi untuk memastikan data tetap akurat dan tidak duplikat',
            ],
            'Level 3 (Established Process)': [
                'Perusahaan memiliki prosedur standar terdokumentasi untuk pengelolaan dan pemutakhiran data',
                'Terdapat kebijakan klasifikasi dan kepemilikan data yang diterapkan secara konsisten',
                'Terdapat mekanisme integrasi data antar saluran sehingga data tersentralisasi',
            ],
        },
    },
    {
        'code': 'BAI09',
        'name': 'Managed Assets (Pengelolaan Aset)',
        'desc': 'Domain ini mengevaluasi pengelolaan aset persediaan barang safety yang menjadi inti bisnis perusahaan.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perusahaan mencatat aset persediaan barang safety yang dimiliki di gudang',
                'Setiap barang masuk dan keluar gudang dicatat untuk mengetahui posisi stok terkini',
                'Perusahaan mengetahui jumlah dan lokasi aset/persediaan yang dimilikinya',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan pengelolaan persediaan (minimum stok, reorder point) yang terstruktur',
                'Posisi stok dipantau dan dilaporkan secara berkala kepada pihak yang berkepentingan',
                'Dilakukan pencocokan (stock opname) antara catatan stok dengan kondisi fisik secara rutin',
            ],
            'Level 3 (Established Process)': [
                'Perusahaan memiliki prosedur standar terdokumentasi untuk pengelolaan persediaan dan aset',
                'Catatan stok terintegrasi dan konsisten antara gudang, toko, dan seluruh marketplace',
                'Terdapat siklus hidup aset yang dikelola dari penerimaan hingga penghapusan secara terdokumentasi',
            ],
        },
    },
    {
        'code': 'DSS01',
        'name': 'Managed Operations (Pengelolaan Operasional)',
        'desc': 'Domain ini mengevaluasi pengelolaan operasional harian, terutama pemenuhan pesanan lintas saluran.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Proses pemenuhan pesanan dari berbagai saluran dijalankan secara rutin setiap hari',
                'Terdapat prosedur penanganan kendala operasional (pesanan salah, stok kosong, keterlambatan)',
                'Perusahaan menjalankan aktivitas operasional harian secara teratur',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan dan penjadwalan operasional yang jelas dan terstruktur',
                'Kinerja operasional dipantau dan dilaporkan secara berkala',
                'Evaluasi operasional dilakukan untuk mengidentifikasi proses yang perlu diperbaiki',
            ],
            'Level 3 (Established Process)': [
                'Perusahaan memiliki SOP operasional yang terdokumentasi untuk seluruh saluran penjualan',
                'SOP operasional diterapkan secara konsisten oleh seluruh staf yang terlibat',
                'Terdapat mekanisme umpan balik dari pelanggan untuk meningkatkan kualitas layanan',
            ],
        },
    },
    {
        'code': 'DSS04',
        'name': 'Managed Continuity (Pengelolaan Keberlangsungan)',
        'desc': 'Domain ini mengevaluasi kesiapan perusahaan menjaga keberlangsungan layanan dan pemulihan saat terjadi gangguan.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perusahaan menyadari adanya risiko gangguan pada saluran penjualan utama (WhatsApp, marketplace, website)',
                'Terdapat upaya pencadangan (backup) data penting seperti data pelanggan dan transaksi',
                'Perusahaan memiliki cara alternatif untuk tetap melayani pelanggan ketika satu saluran terganggu',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan keberlangsungan bisnis (business continuity) yang disusun secara terstruktur',
                'Proses pencadangan data dilakukan secara terjadwal dan dipantau keberhasilannya',
                'Kesiapan menghadapi gangguan dievaluasi secara berkala',
            ],
            'Level 3 (Established Process)': [
                'Perusahaan memiliki dokumen rencana keberlangsungan dan pemulihan (BCP/DRP) yang terdokumentasi formal',
                'Prosedur pemulihan diuji coba secara berkala untuk memastikan dapat berjalan saat dibutuhkan',
                'Terdapat penetapan target waktu pemulihan (RTO) dan titik pemulihan data (RPO) yang disepakati',
            ],
        },
    },
    {
        'code': 'DSS05',
        'name': 'Managed Security Services (Pengelolaan Layanan Keamanan)',
        'desc': 'Domain ini mengevaluasi implementasi teknis pengamanan akun, data, dan perangkat perusahaan.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perangkat yang digunakan perusahaan dilindungi dari malware dan virus',
                'Terdapat pengaturan hak akses ke akun marketplace, data pelanggan, dan keuangan',
                'Akun penting perusahaan dilindungi dengan kata sandi dan dijaga kerahasiaannya',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan yang terstruktur dalam mengelola keamanan akun dan data',
                'Aktivitas akses ke sistem dan akun penting dipantau serta insiden keamanan dilaporkan',
                'Efektivitas pengamanan akun dan data dievaluasi secara berkala',
            ],
            'Level 3 (Established Process)': [
                'Perusahaan memiliki prosedur standar keamanan informasi yang terdokumentasi',
                'Pengaturan hak akses ditinjau secara berkala saat karyawan masuk, pindah peran, atau keluar',
                'Terdapat pemeriksaan keamanan internal secara berkala untuk memastikan kepatuhan',
            ],
        },
    },
]

q_global = 1
for d_idx, domain in enumerate(domains):
    centered(f"DOMAIN {domain['code']}", 14, bold=True, space_after=2)
    centered(domain['name'], 12, bold=True, space_after=6)
    normal(domain['desc'], space_after=12)

    for level_name, questions in domain['levels'].items():
        normal(level_name, bold=True, space_after=6)
        num_q = len(questions)
        table = doc.add_table(rows=1 + num_q, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(["No", "Pernyataan", "STS (1)", "TS (2)", "R (3)", "S (4)", "SS (5)"]):
            cell = table.rows[0].cells[j]
            cell.text = h
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_FILL}"/>'))
        for qi, question in enumerate(questions):
            row = table.rows[qi + 1]
            row.cells[0].text = str(q_global)
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in row.cells[0].paragraphs[0].runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            row.cells[1].text = question
            for r in row.cells[1].paragraphs[0].runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            for ci in range(2, 7):
                row.cells[ci].text = ""
                row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row.cells[ci].width = Cm(1.5)
            row.cells[0].width = Cm(1)
            row.cells[1].width = Cm(8)
            q_global += 1
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    if d_idx < len(domains) - 1:
        doc.add_page_break()

doc.add_paragraph()
normal("— Terima kasih atas kesediaan Bapak/Ibu mengisi kuesioner ini —", bold=True, space_after=6)
normal("Data yang Bapak/Ibu berikan akan dijaga kerahasiaannya dan hanya digunakan untuk kepentingan akademis.", space_after=6)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(24)
r = p.add_run("Tanda tangan responden:"); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("(__________________________)"); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

doc.save(OUTPUT)
print(f"OK Form Kuesioner: {OUTPUT}")
print(f"Total pertanyaan: {q_global - 1} (9 per domain x 5 domain = 45)")
