"""
Generate Form Kuesioner TERISI per Responden — Audit SI PT Aegisindo Mitra Sejati
Menghasilkan 6 file DOCX (1 per responden) dengan jawaban sudah terisi (centang).
Data jawaban identik dengan Tabulasi_Data.md.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "form_responden")
os.makedirs(OUTPUT_DIR, exist_ok=True)

HEADER_FILL = "F26522"  # safety orange


# ============================================================
# DATA RESPONDEN (konsisten dengan Tabulasi_Data.md & Master Excel)
# ============================================================
RESPONDENTS = [
    {'code': 'R1', 'initial': 'MO', 'name': 'Mada Oktavian',
     'jabatan': 'Manajer Operasional', 'divisi': 'Operasional',
     'lama': '7 tahun', 'pendidikan': 'S1 Manajemen', 'tanggal': '10 / 03 / 2026'},
    {'code': 'R2', 'initial': 'PG', 'name': 'Panji Gunawan',
     'jabatan': 'Staff Procurement & Gudang', 'divisi': 'Pengadaan & Gudang',
     'lama': '5 tahun', 'pendidikan': 'S1 Teknik Industri', 'tanggal': '10 / 03 / 2026'},
    {'code': 'R3', 'initial': 'PO', 'name': 'Putri Oktaviani',
     'jabatan': 'Staff Penjualan Online (Marketplace)', 'divisi': 'Penjualan & Marketing',
     'lama': '4 tahun', 'pendidikan': 'S1 Manajemen', 'tanggal': '11 / 03 / 2026'},
    {'code': 'R4', 'initial': 'AS', 'name': 'Andi Saputra',
     'jabatan': 'Staff IT / Admin Sistem', 'divisi': 'IT / Administrasi',
     'lama': '3 tahun', 'pendidikan': 'S1 Sistem Informasi', 'tanggal': '11 / 03 / 2026'},
    {'code': 'R5', 'initial': 'KA', 'name': 'Kartika Anggraini',
     'jabatan': 'Staff Keuangan & Administrasi', 'divisi': 'Keuangan & Administrasi',
     'lama': '5 tahun', 'pendidikan': 'S1 Akuntansi', 'tanggal': '12 / 03 / 2026'},
    {'code': 'R6', 'initial': 'CS', 'name': 'Cahya Saputri',
     'jabatan': 'Staff Customer Service', 'divisi': 'Customer Service & After Sales',
     'lama': '2 tahun', 'pendidikan': 'D3 Administrasi Bisnis', 'tanggal': '12 / 03 / 2026'},
]

# Jawaban per domain — 9 pernyataan (L1: q1-3, L2: q4-6, L3: q7-9), nilai [R1..R6]
ANSWERS = {
    'APO14': [
        [4, 3, 3, 4, 3, 2], [4, 3, 3, 3, 3, 2], [3, 3, 3, 3, 3, 2],
        [3, 3, 2, 3, 2, 2], [3, 2, 2, 2, 2, 2], [3, 2, 2, 2, 2, 2],
        [2, 2, 2, 2, 2, 1], [2, 2, 2, 2, 2, 1], [2, 2, 2, 2, 1, 1],
    ],
    'BAI09': [
        [4, 4, 3, 4, 3, 3], [4, 4, 3, 4, 3, 3], [4, 3, 3, 3, 3, 3],
        [3, 3, 3, 3, 3, 2], [3, 3, 3, 3, 2, 2], [3, 3, 2, 3, 2, 2],
        [3, 2, 2, 3, 2, 2], [3, 2, 2, 2, 2, 2], [2, 2, 2, 2, 2, 2],
    ],
    'DSS01': [
        [4, 4, 4, 4, 3, 3], [4, 4, 3, 4, 3, 3], [4, 3, 3, 4, 3, 3],
        [4, 3, 3, 3, 3, 2], [3, 3, 3, 3, 3, 2], [3, 3, 3, 3, 2, 2],
        [3, 3, 2, 3, 2, 2], [3, 2, 2, 3, 2, 2], [3, 2, 2, 2, 2, 2],
    ],
    'DSS04': [
        [4, 3, 3, 3, 3, 2], [3, 3, 3, 3, 3, 2], [3, 3, 3, 3, 3, 2],
        [3, 2, 2, 3, 2, 2], [3, 2, 2, 2, 2, 2], [2, 2, 2, 2, 2, 2],
        [2, 2, 2, 2, 2, 1], [2, 2, 2, 2, 1, 1], [2, 2, 1, 2, 1, 1],
    ],
    'DSS05': [
        [4, 4, 3, 4, 3, 3], [4, 3, 3, 4, 3, 3], [3, 3, 3, 3, 3, 3],
        [3, 3, 3, 3, 3, 2], [3, 3, 2, 3, 2, 2], [3, 3, 2, 3, 2, 2],
        [3, 2, 2, 2, 2, 2], [2, 2, 2, 2, 2, 2], [2, 2, 2, 2, 2, 2],
    ],
}

# Domain + pernyataan (identik dengan generate_form_kuesioner.py)
DOMAINS = [
    {
        'code': 'APO14', 'name': 'Managed Data (Pengelolaan Data)',
        'desc': 'Domain ini mengevaluasi sejauh mana perusahaan mengelola data inventori, pelanggan, dan penjualan secara terstruktur.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perusahaan telah mendata dan mengelola data inventori, pelanggan, dan penjualan yang dibutuhkan operasional',
                'Data produk dan stok dicatat dan diperbarui saat terjadi transaksi pembelian maupun penjualan',
                'Perusahaan menyimpan data pelanggan dan riwayat transaksi pada media penyimpanan tertentu',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan yang jelas mengenai data apa saja yang dikelola dan siapa penanggung jawabnya',
                'Kualitas dan konsistensi data antar saluran penjualan dipantau secara berkala',
                'Hasil pengelolaan data dievaluasi untuk memastikan data tetap akurat dan tidak duplikat',
            ],
            'Level 3 (Established Process)': [
                'Perusahaan memiliki prosedur standar terdokumentasi untuk pengelolaan dan pemutakhiran data',
                'Terdapat kebijakan klasifikasi dan kepemilikan data yang diterapkan secara konsisten',
                'Terdapat mekanisme integrasi data antar saluran sehingga data tersentralisasi',
            ],
        },
    },
    {
        'code': 'BAI09', 'name': 'Managed Assets (Pengelolaan Aset)',
        'desc': 'Domain ini mengevaluasi pengelolaan aset persediaan barang safety yang menjadi inti bisnis perusahaan.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perusahaan mencatat aset persediaan barang safety yang dimiliki di gudang',
                'Setiap barang masuk dan keluar gudang dicatat untuk mengetahui posisi stok terkini',
                'Perusahaan mengetahui jumlah dan lokasi aset/persediaan yang dimilikinya',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan pengelolaan persediaan (minimum stok, reorder point) yang terstruktur',
                'Posisi stok dipantau dan dilaporkan secara berkala kepada pihak yang berkepentingan',
                'Dilakukan pencocokan (stock opname) antara catatan stok dengan kondisi fisik secara rutin',
            ],
            'Level 3 (Established Process)': [
                'Perusahaan memiliki prosedur standar terdokumentasi untuk pengelolaan persediaan dan aset',
                'Catatan stok terintegrasi dan konsisten antara gudang, toko, dan seluruh marketplace',
                'Terdapat siklus hidup aset yang dikelola dari penerimaan hingga penghapusan secara terdokumentasi',
            ],
        },
    },
    {
        'code': 'DSS01', 'name': 'Managed Operations (Pengelolaan Operasional)',
        'desc': 'Domain ini mengevaluasi pengelolaan operasional harian, terutama pemenuhan pesanan lintas saluran.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Proses pemenuhan pesanan dari berbagai saluran dijalankan secara rutin setiap hari',
                'Terdapat prosedur penanganan kendala operasional (pesanan salah, stok kosong, keterlambatan)',
                'Perusahaan menjalankan aktivitas operasional harian secara teratur',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan dan penjadwalan operasional yang jelas dan terstruktur',
                'Kinerja operasional dipantau dan dilaporkan secara berkala',
                'Evaluasi operasional dilakukan untuk mengidentifikasi proses yang perlu diperbaiki',
            ],
            'Level 3 (Established Process)': [
                'Perusahaan memiliki SOP operasional yang terdokumentasi untuk seluruh saluran penjualan',
                'SOP operasional diterapkan secara konsisten oleh seluruh staf yang terlibat',
                'Terdapat mekanisme umpan balik dari pelanggan untuk meningkatkan kualitas layanan',
            ],
        },
    },
    {
        'code': 'DSS04', 'name': 'Managed Continuity (Pengelolaan Keberlangsungan)',
        'desc': 'Domain ini mengevaluasi kesiapan perusahaan menjaga keberlangsungan layanan dan pemulihan saat terjadi gangguan.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perusahaan menyadari adanya risiko gangguan pada saluran penjualan utama (WhatsApp, marketplace, website)',
                'Terdapat upaya pencadangan (backup) data penting seperti data pelanggan dan transaksi',
                'Perusahaan memiliki cara alternatif untuk tetap melayani pelanggan ketika satu saluran terganggu',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan keberlangsungan bisnis (business continuity) yang disusun secara terstruktur',
                'Proses pencadangan data dilakukan secara terjadwal dan dipantau keberhasilannya',
                'Kesiapan menghadapi gangguan dievaluasi secara berkala',
            ],
            'Level 3 (Established Process)': [
                'Perusahaan memiliki dokumen rencana keberlangsungan dan pemulihan (BCP/DRP) yang terdokumentasi formal',
                'Prosedur pemulihan diuji coba secara berkala untuk memastikan dapat berjalan saat dibutuhkan',
                'Terdapat penetapan target waktu pemulihan (RTO) dan titik pemulihan data (RPO) yang disepakati',
            ],
        },
    },
    {
        'code': 'DSS05', 'name': 'Managed Security Services (Pengelolaan Layanan Keamanan)',
        'desc': 'Domain ini mengevaluasi implementasi teknis pengamanan akun, data, dan perangkat perusahaan.',
        'levels': {
            'Level 1 (Performed Process)': [
                'Perangkat yang digunakan perusahaan dilindungi dari malware dan virus',
                'Terdapat pengaturan hak akses ke akun marketplace, data pelanggan, dan keuangan',
                'Akun penting perusahaan dilindungi dengan kata sandi dan dijaga kerahasiaannya',
            ],
            'Level 2 (Managed Process)': [
                'Terdapat perencanaan yang terstruktur dalam mengelola keamanan akun dan data',
                'Aktivitas akses ke sistem dan akun penting dipantau serta insiden keamanan dilaporkan',
                'Efektivitas pengamanan akun dan data dievaluasi secara berkala',
            ],
            'Level 3 (Established Process)': [
                'Perusahaan memiliki prosedur standar keamanan informasi yang terdokumentasi',
                'Pengaturan hak akses ditinjau secara berkala saat karyawan masuk, pindah peran, atau keluar',
                'Terdapat pemeriksaan keamanan internal secara berkala untuk memastikan kepatuhan',
            ],
        },
    },
]


def generate_form(resp_idx, resp):
    """Generate satu DOCX form terisi untuk satu responden."""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)

    def centered(text, size=12, bold=False, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold

    def normal(text, bold=False, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = bold

    def set_cell_font(cell, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, bg=None):
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = align
        for r in p.runs:
            r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold
            if color:
                r.font.color.rgb = color
        if bg:
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>'))

    # ── COVER ──
    centered("KUESIONER PENELITIAN", 16, bold=True, space_after=12)
    centered("AUDIT SISTEM INFORMASI", 14, bold=True, space_after=6)
    centered("PT AEGISINDO MITRA SEJATI", 14, bold=True, space_after=6)
    centered("Framework COBIT 2019", 12, space_after=18)
    centered(f"— Form Isian Responden {resp['code']} —", 12, bold=True, space_after=24)

    # ── IDENTITAS RESPONDEN (TERISI) ──
    centered("IDENTITAS RESPONDEN", 14, bold=True, space_after=12)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    fields = [
        ("Nama", resp['name']),
        ("Jabatan", resp['jabatan']),
        ("Divisi / Bagian", resp['divisi']),
        ("Lama Bekerja", resp['lama']),
        ("Pendidikan Terakhir", resp['pendidikan']),
        ("Tanggal Pengisian", resp['tanggal']),
    ]
    for i, (label, value) in enumerate(fields):
        set_cell_font(table.rows[i].cells[0], label, 11)
        set_cell_font(table.rows[i].cells[1], ":", 11, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_font(table.rows[i].cells[2], value, 11, bold=True)
        table.rows[i].cells[0].width = Cm(5)
        table.rows[i].cells[1].width = Cm(0.5)
        table.rows[i].cells[2].width = Cm(9)

    doc.add_page_break()

    # ── KUESIONER PER DOMAIN (JAWABAN TERISI) ──
    q_global = 0
    for d_idx, domain in enumerate(DOMAINS):
        centered(f"DOMAIN {domain['code']}", 14, bold=True, space_after=2)
        centered(domain['name'], 12, bold=True, space_after=6)
        normal(domain['desc'], space_after=12)

        for level_name, questions in domain['levels'].items():
            normal(level_name, bold=True, space_after=6)
            num_q = len(questions)
            table = doc.add_table(rows=1 + num_q, cols=7)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for j, h in enumerate(["No", "Pernyataan", "STS (1)", "TS (2)", "R (3)", "S (4)", "SS (5)"]):
                set_cell_font(table.rows[0].cells[j], h, 10, bold=True,
                              align=WD_ALIGN_PARAGRAPH.CENTER,
                              color=RGBColor(255, 255, 255), bg=HEADER_FILL)

            for qi, question in enumerate(questions):
                row = table.rows[qi + 1]
                set_cell_font(row.cells[0], str(q_global + 1), 10, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_font(row.cells[1], question, 10)

                answer = ANSWERS[domain['code']][q_global % 9][resp_idx]
                for ci in range(2, 7):
                    score = ci - 1  # kolom 2=STS(1) ... kolom 6=SS(5)
                    if score == answer:
                        set_cell_font(row.cells[ci], "✓", 12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                    else:
                        row.cells[ci].text = ""
                        row.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row.cells[ci].width = Cm(1.5)
                row.cells[0].width = Cm(1)
                row.cells[1].width = Cm(8)
                q_global += 1

            doc.add_paragraph().paragraph_format.space_after = Pt(8)

        if d_idx < len(DOMAINS) - 1:
            doc.add_page_break()

    # ── FOOTER ──
    doc.add_paragraph()
    normal("— Terima kasih atas kesediaan Bapak/Ibu mengisi kuesioner ini —", bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run("Tanda tangan responden:"); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"( {resp['name']} )"); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.font.bold = True

    filename = f"Kuesioner_Terisi_{resp['code']}_{resp['initial']}.docx"
    doc.save(os.path.join(OUTPUT_DIR, filename))
    return filename


if __name__ == '__main__':
    print("=" * 60)
    print("  GENERATE FORM KUESIONER TERISI PER RESPONDEN")
    print("=" * 60)
    for idx, resp in enumerate(RESPONDENTS):
        fname = generate_form(idx, resp)
        print(f"  OK {resp['code']} ({resp['initial']}) - {resp['jabatan']:38s} -> {fname}")
    print(f"\nOutput folder: {OUTPUT_DIR}")
    print(f"Total file: {len(RESPONDENTS)} form kuesioner terisi (45 jawaban centang/file)")
