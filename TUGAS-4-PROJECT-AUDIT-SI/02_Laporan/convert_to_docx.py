"""
Konversi Laporan Project Audit SI PT Aegisindo Mitra Sejati dari Markdown ke Word (.docx)
Format Akademis: Margin 4-3-2.5-2.5, TNR 12pt, Spasi 1.5
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

LAPORAN_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(LAPORAN_DIR)
CHARTS_DIR = os.path.join(BASE_DIR, "03_Data", "charts")
LOGO_PATH = os.path.join(BASE_DIR, "05_Presentasi", "assets", "images", "logo-unm.png")
OUTPUT_PATH = os.path.join(LAPORAN_DIR, "Kelompok-Laporan-Audit-SI-PT_Aegisindo_Mitra_Sejati.docx")

# === GAMBAR MAP: caption keyword -> (file, width_cm) ===
GAMBAR_MAP = {
    'Gambar 1.1': ('gambar_5_struktur_organisasi.png', 14),
    'Gambar 2.1': ('gambar_4_metodologi.png', 14),
    'Gambar 2.2': ('gambar_1_radar_capability.png', 13),
    'Gambar 2.3': ('gambar_2_gap_analysis.png', 14),
    'Gambar 2.4': ('gambar_3_ketercapaian.png', 14),
}

doc = Document()

def setup_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(4)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.footer_distance = Cm(0.3)
    section.header_distance = Cm(1.25)

for section in doc.sections:
    setup_section(section)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.line_spacing = 1.5
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
doc.styles['Heading 1'].font.size = Pt(14)
doc.styles['Heading 2'].font.size = Pt(12)
doc.styles['Heading 3'].font.size = Pt(12)


# === PAGINATION HELPERS ===
def add_page_field(paragraph, fmt="arabic"):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin'); run._r.append(fldChar1)
    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE \\* roman ' if fmt == "roman" else ' PAGE '
    run2._r.append(instrText)
    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end'); run3._r.append(fldChar2)
    for r in paragraph.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)

def set_page_number_start(section, start=1, fmt="decimal"):
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), str(start)); pgNumType.set(qn('w:fmt'), fmt)
    section._sectPr.append(pgNumType)

def clear_header_footer(section):
    for attr in ['header', 'footer', 'first_page_header', 'first_page_footer']:
        try:
            obj = getattr(section, attr)
            obj.is_linked_to_previous = False
            for p in obj.paragraphs:
                p.text = ''
        except Exception:
            pass

def setup_cover_section(section):
    setup_section(section); clear_header_footer(section)

def setup_front_matter_section(section, start_page=2):
    setup_section(section)
    set_page_number_start(section, start=start_page, fmt="lowerRoman")
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp, fmt="roman")
    section.header.is_linked_to_previous = False
    for p in section.header.paragraphs:
        p.text = ''

def setup_bab_section(section, restart_at=None):
    setup_section(section)
    section.different_first_page_header_footer = True
    sectPr = section._sectPr
    if restart_at is not None:
        set_page_number_start(section, start=restart_at, fmt="decimal")
    else:
        for pgNum in sectPr.findall(qn('w:pgNumType')):
            sectPr.remove(pgNum)
    section.first_page_footer.is_linked_to_previous = False
    fp = section.first_page_footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp, fmt="arabic")
    section.first_page_header.is_linked_to_previous = False
    for p in section.first_page_header.paragraphs:
        p.text = ''
    section.header.is_linked_to_previous = False
    hp = section.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(hp, fmt="arabic")
    section.footer.is_linked_to_previous = False
    for p in section.footer.paragraphs:
        p.text = ''

def setup_plain_arabic_section(section):
    setup_section(section)
    section.different_first_page_header_footer = False
    sectPr = section._sectPr
    for pgNum in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(pgNum)
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp, fmt="arabic")
    section.header.is_linked_to_previous = False
    for p in section.header.paragraphs:
        p.text = ''

def add_new_section(section_type=WD_SECTION.NEW_PAGE):
    doc.add_section(section_type)
    return doc.sections[-1]

def add_toc_entry(text, page_text, is_bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = is_bold
    r2 = p.add_run(f"\t{page_text}"); r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)


# === CONTENT HELPERS ===
def centered(text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.font.bold = bold
    return p

def normal(text, bold=False, italic=False, indent=True):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True
        else:
            italic_parts = re.split(r'(\*[^*]+?\*)', part)
            for ipart in italic_parts:
                if ipart.startswith('*') and ipart.endswith('*') and not ipart.startswith('**'):
                    r = p.add_run(ipart[1:-1]); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.italic = True
                else:
                    r = p.add_run(ipart); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
                    r.font.bold = bold; r.font.italic = italic
    return p

def strip_md(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text

def list_item(text, level=0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.25 + level * 0.75)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    parts = re.split(r'(\*\*.*?\*\*)', strip_md(text))
    for part in parts:
        r = p.add_run(part); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit = True
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), '5000'); tblW.set(qn('w:type'), 'pct')
    table._tbl.tblPr.append(tblW)
    for c_idx, h in enumerate(headers):
        if h.strip().lower() in ["no", "no."]:
            table.columns[c_idx].width = Cm(1.2)
    hdr_row = table.add_row()
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]; cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(h); r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.bold = True
        cp.paragraph_format.line_spacing = 1.0; cp.paragraph_format.space_after = Pt(2); cp.paragraph_format.space_before = Pt(2)
    for row_data in rows:
        row = table.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]; cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = cp.add_run(strip_md(str(val))); r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            cp.paragraph_format.line_spacing = 1.0; cp.paragraph_format.space_after = Pt(2); cp.paragraph_format.space_before = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def add_image(image_path, caption, width_cm=14):
    if not os.path.exists(image_path):
        print(f"[WARN] Gambar tidak ditemukan: {image_path}")
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(image_path, width=Cm(width_cm))
    p_cap = doc.add_paragraph(); p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    r_cap = p_cap.add_run(caption); r_cap.font.name = 'Times New Roman'; r_cap.font.size = Pt(10); r_cap.font.italic = True

def parse_md_table(text):
    lines = [l.strip() for l in text.strip().split('\n') if l.strip()]
    lines = [l for l in lines if not re.match(r'^\|[-\s|:]+\|$', l)]
    if not lines:
        return [], []
    headers = [c.strip() for c in lines[0].split('|')[1:-1]]
    rows = [[c.strip() for c in l.split('|')[1:-1]] for l in lines[1:]]
    return headers, [r for r in rows if r]

def process_md_content(content, skip_title=True):
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line or line.strip() == '---':
            i += 1; continue
        if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
            i += 1; continue
        if re.match(r'^[-:|\s]+$', line.strip()) and len(line.strip()) <= 10:
            i += 1; continue
        if skip_title and (line.startswith('# BAB') or line.startswith('# DAFTAR')):
            i += 1; continue
        if line.startswith('#### '):
            p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(strip_md(line[5:].strip())); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True
            i += 1; continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3); i += 1; continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2); i += 1; continue
        bold_match = re.match(r'^\*\*(.+?)\*\*$', line.strip())
        if bold_match and not line.strip().startswith('|') and line.strip().count('**') == 2:
            text = bold_match.group(1)
            p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
            if text.startswith('Tabel') or text.startswith('Gambar'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
            else:
                p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True
            i += 1; continue
        if line.strip().startswith('|'):
            table_lines = [line]; i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i]); i += 1
            headers, rows = parse_md_table('\n'.join(table_lines))
            if headers and rows:
                add_table(headers, rows)
            continue
        if line.strip().startswith('>'):
            bq_text = line.strip().lstrip('>').strip()
            gambar_match = re.match(r'^\*\*(.+?)\*\*\s*(.*)', bq_text)
            if gambar_match:
                gambar_key = gambar_match.group(1).strip()
                gambar_caption = gambar_match.group(2).strip()
                full_caption = f"{gambar_key} — {gambar_caption}" if gambar_caption else gambar_key
                for key, (filename, width) in GAMBAR_MAP.items():
                    if key in gambar_key:
                        add_image(os.path.join(CHARTS_DIR, filename), full_caption, width_cm=width)
                        break
            else:
                normal(strip_md(bq_text), italic=True)
            i += 1; continue
        if re.match(r'^[a-z]\.\s', line.strip()):
            list_item(line.strip(), level=1); i += 1; continue
        if re.match(r'^\d+\.', line.strip()):
            list_item(line.strip()); i += 1; continue
        if line.strip().startswith('- ') or line.strip().startswith('• '):
            text = line.strip().lstrip('-•').strip()
            list_item(f"• {strip_md(text)}"); i += 1; continue
        text = line.strip()
        if text and not text.startswith(('#', '|', '!', '>')):
            normal(text)
        i += 1


# ============================================================
# SECTION 1: COVER
# ============================================================
print(">> Cover...")
setup_cover_section(doc.sections[0])
centered("LAPORAN PROJECT", 16, bold=True, space_after=0)
centered("AUDIT SISTEM INFORMASI", 16, bold=True, space_after=18)
if os.path.exists(LOGO_PATH):
    p_logo = doc.add_paragraph(); p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(18); p_logo.paragraph_format.line_spacing = 1.0
    p_logo.add_run().add_picture(LOGO_PATH, width=Cm(5), height=Cm(5))
centered("AUDIT TATA KELOLA TEKNOLOGI INFORMASI", 14, bold=True, space_after=0)
centered("MENGGUNAKAN FRAMEWORK COBIT 2019", 14, bold=True, space_after=0)
centered("PADA PT AEGISINDO MITRA SEJATI", 14, bold=True, space_after=18)
centered("Disusun untuk Memenuhi Tugas Mata Kuliah", 12, space_after=0)
centered("Audit Sistem Informasi (453)", 12, bold=True, space_after=18)
centered("Disusun Oleh:", 12, bold=True, space_after=8)
centered("Dinda Callista (11250083)", 12, space_after=0)
centered("Alim Fatun Rofiah (11250127)", 12, space_after=0)
centered("Nur Fikri (11250138)", 12, space_after=18)
centered("Dosen Pengampu:", 12, bold=True, space_after=0)
centered("Juarni Siregar, S.Pd., M.Kom", 12, bold=True, space_after=18)
centered("PROGRAM STUDI SISTEM INFORMASI (S1)", 12, bold=True, space_after=0)
centered("FAKULTAS TEKNOLOGI INFORMASI", 12, bold=True, space_after=0)
centered("UNIVERSITAS NUSA MANDIRI", 12, bold=True, space_after=0)
centered("JAKARTA", 12, bold=True, space_after=0)
centered("2026", 12, bold=True)


# ============================================================
# SECTION 2: FRONT MATTER (romawi)
# ============================================================
print(">> Front Matter...")
setup_front_matter_section(add_new_section(), start_page=2)
centered("KATA PENGANTAR", 14, bold=True, space_after=24)
for para in [
    "Puji syukur kami panjatkan kehadirat Tuhan Yang Maha Esa, karena atas berkat dan rahmat-Nya kami dapat menyelesaikan laporan project audit sistem informasi ini dengan baik. Laporan ini disusun untuk memenuhi tugas mata kuliah Audit Sistem Informasi (453) pada Program Studi Sistem Informasi (S1), Fakultas Teknologi Informasi, Universitas Nusa Mandiri.",
    "Laporan ini membahas audit tata kelola teknologi informasi pada PT Aegisindo Mitra Sejati menggunakan framework COBIT 2019. Melalui audit ini, kami melakukan penilaian terhadap tingkat kapabilitas (capability level) tata kelola TI pada lima domain yang relevan dengan kondisi perusahaan, yaitu APO14 (Managed Data), BAI09 (Managed Assets), DSS01 (Managed Operations), DSS04 (Managed Continuity), dan DSS05 (Managed Security Services).",
    "Kami menyadari bahwa laporan ini masih jauh dari sempurna. Oleh karena itu, kami mengharapkan kritik dan saran yang membangun dari Ibu dosen pengampu dan rekan-rekan mahasiswa guna penyempurnaan laporan ini di kemudian hari.",
    "Kami mengucapkan terima kasih kepada Ibu Juarni Siregar, S.Pd., M.Kom selaku dosen pengampu, PT Aegisindo Mitra Sejati sebagai objek studi kasus, seluruh responden yang telah berpartisipasi, serta semua pihak yang telah membantu penyelesaian laporan ini.",
    "Semoga laporan ini dapat memberikan manfaat dan kontribusi dalam pengembangan tata kelola teknologi informasi, khususnya bagi perusahaan distribusi alat keselamatan kerja.",
]:
    normal(para)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.paragraph_format.space_before = Pt(24)
r = p.add_run("Jakarta, Mei 2026"); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Tim Penulis"); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.bold = True
doc.add_page_break()

# DAFTAR ISI
centered("DAFTAR ISI", 14, bold=True, space_after=24)
toc = [
    ("HALAMAN JUDUL", "i", True), ("KATA PENGANTAR", "ii", True), ("DAFTAR ISI", "iii", True),
    ("DAFTAR TABEL", "iv", True), ("DAFTAR GAMBAR", "v", True), ("", "", False),
    ("BAB I PENDAHULUAN", "1", True), ("A. Latar Belakang", "1", False),
    ("B. Permasalahan", "2", False), ("C. Tujuan", "3", False), ("", "", False),
    ("BAB II PEMBAHASAN", "4", True), ("A. Framework yang Digunakan", "4", False),
    ("B. Domain dan Subdomain yang Digunakan", "6", False), ("C. Capability Level dan GAP", "10", False),
    ("", "", False), ("BAB III PENUTUP", "18", True), ("A. Kesimpulan", "18", False),
    ("B. Saran", "19", False), ("", "", False), ("DAFTAR PUSTAKA", "21", True),
    ("LAMPIRAN", "23", True), ("Lampiran 1: Kuesioner Penelitian", "23", False),
    ("Lampiran 2: Tabulasi Data Responden", "33", False), ("Lampiran 3: Perhitungan Capability Level", "42", False),
]
for item, page, is_bold in toc:
    if not item:
        doc.add_paragraph(); continue
    add_toc_entry(item, page, is_bold=is_bold, indent=not is_bold)
doc.add_page_break()

# DAFTAR TABEL
centered("DAFTAR TABEL", 14, bold=True, space_after=24)
tables_list = [
    ("Tabel 2.1", "Perbedaan COBIT 5 dan COBIT 2019"),
    ("Tabel 2.2", "Capability Level COBIT 2019 (Level 0–5)"),
    ("Tabel 2.3", "Rating Scale Ketercapaian"),
    ("Tabel 2.4", "RACI Matrix"),
    ("Tabel 2.5", "Hasil Perhitungan Capability Level APO14"),
    ("Tabel 2.6", "Hasil Perhitungan Capability Level BAI09"),
    ("Tabel 2.7", "Hasil Perhitungan Capability Level DSS01"),
    ("Tabel 2.8", "Hasil Perhitungan Capability Level DSS04"),
    ("Tabel 2.9", "Hasil Perhitungan Capability Level DSS05"),
    ("Tabel 2.10", "Ringkasan Capability Level dan GAP Analysis"),
]
for num, title in tables_list:
    add_toc_entry(f"{num} — {title}", "...")
doc.add_page_break()

# DAFTAR GAMBAR
centered("DAFTAR GAMBAR", 14, bold=True, space_after=24)
gambar_list = [
    ("Gambar 1.1", "Peta Fungsi Organisasi dan Fokus Audit PT Aegisindo"),
    ("Gambar 2.1", "Tahapan Audit Sistem Informasi PT Aegisindo"),
    ("Gambar 2.2", "Profil Capability Proses TI PT Aegisindo"),
    ("Gambar 2.3", "Peta Kesenjangan As-is dan To-be per Domain"),
    ("Gambar 2.4", "Matriks Rating Ketercapaian Capability per Domain"),
]
for num, title in gambar_list:
    add_toc_entry(f"{num} — {title}", "...")


# ============================================================
# SECTION 3-5: BAB I, II, III
# ============================================================
print(">> BAB I...")
setup_bab_section(add_new_section(), restart_at=1)
centered("BAB I", 14, bold=True, space_after=0)
centered("PENDAHULUAN", 14, bold=True, space_after=24)
process_md_content(open(os.path.join(LAPORAN_DIR, "BAB-I_Pendahuluan.md"), encoding='utf-8').read())

print(">> BAB II...")
setup_bab_section(add_new_section())
centered("BAB II", 14, bold=True, space_after=0)
centered("PEMBAHASAN", 14, bold=True, space_after=24)
process_md_content(open(os.path.join(LAPORAN_DIR, "BAB-II_Pembahasan.md"), encoding='utf-8').read())

print(">> BAB III...")
setup_bab_section(add_new_section())
centered("BAB III", 14, bold=True, space_after=0)
centered("PENUTUP", 14, bold=True, space_after=24)
process_md_content(open(os.path.join(LAPORAN_DIR, "BAB-III_Penutup.md"), encoding='utf-8').read())


# ============================================================
# SECTION 6: DAFTAR PUSTAKA
# ============================================================
print(">> Daftar Pustaka...")
setup_plain_arabic_section(add_new_section())
centered("DAFTAR PUSTAKA", 14, bold=True, space_after=24)
for line in open(os.path.join(LAPORAN_DIR, "Daftar_Pustaka.md"), encoding='utf-8').read().split('\n'):
    line = line.strip()
    if not line or line.startswith('#'):
        continue
    text = re.sub(r'\*(.+?)\*', r'\1', line)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.25); p.paragraph_format.first_line_indent = Cm(-1.25)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12)


# ============================================================
# SECTION 7: LAMPIRAN
# ============================================================
print(">> Lampiran...")
DATA_DIR = os.path.join(BASE_DIR, "03_Data")

def strip_md_title_block(text):
    lines = text.split('\n'); out = []; skipping = True
    for ln in lines:
        s = ln.strip()
        if skipping:
            if s == '' or s.startswith('# ') or s.startswith('## ') or s == '---':
                continue
            skipping = False
        out.append(ln)
    return '\n'.join(out)

lampiran_items = [
    ("Lampiran 1", "Kuesioner Penelitian", "Kuesioner.md",
     "Lampiran ini menyajikan instrumen kuesioner (dalam keadaan kosong) yang dibagikan kepada responden untuk mengumpulkan data penilaian terhadap setiap pernyataan. Rekapitulasi jawaban keenam responden disajikan pada Lampiran 2, sedangkan RACI Matrix di bawah merupakan pemetaan tanggung jawab peran dan bukan bagian yang diisi oleh responden."),
    ("Lampiran 2", "Tabulasi Data Responden", "Tabulasi_Data.md",
     "Lampiran ini menyajikan rekapitulasi jawaban kuesioner dari keenam responden (R1–R6) beserta nilai total dan rata-rata (mean) untuk setiap pernyataan, yang menjadi dasar perhitungan capability level pada Lampiran 3."),
    ("Lampiran 3", "Perhitungan Capability Level", "Perhitungan_Capability_Level.md",
     "Lampiran ini menyajikan rincian perhitungan capability level setiap domain COBIT 2019 berdasarkan data pada Lampiran 2, mulai dari nilai mean per level, persentase ketercapaian, rating (N/P/L/F), hingga penentuan capability level dan analisis GAP."),
]
setup_plain_arabic_section(add_new_section())
centered("LAMPIRAN", 16, bold=True, space_after=18)
for idx, (label, title, filename, keterangan) in enumerate(lampiran_items):
    if idx > 0:
        doc.add_page_break()
    centered(label, 14, bold=True, space_after=0)
    centered(title, 14, bold=True, space_after=18)
    if keterangan:
        normal(keterangan, italic=True)
    content = strip_md_title_block(open(os.path.join(DATA_DIR, filename), encoding='utf-8').read())
    process_md_content(content, skip_title=False)


doc.save(OUTPUT_PATH)
print(f"[OK] Laporan Word berhasil disimpan: {OUTPUT_PATH}")
print(f"[i] Format: TNR 12pt, Spasi 1.5, Margin 4-3-2.5-2.5 cm")
